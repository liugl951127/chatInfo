# -*- coding: utf-8 -*-
"""
build_op_manual.py - 生成 OPERATION-MANUAL.docx 操作手册
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

d = Document()
style = d.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for s in d.styles:
    if s.name.startswith('Heading'):
        s.font.name = 'Microsoft YaHei'
        s.font.color.rgb = RGBColor(0x1e, 0x40, 0x80)

# 封面
title = d.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('\n\n\nV3.0 智能客服平台')
run.font.size = Pt(36); run.font.bold = True
run.font.color.rgb = RGBColor(0x1e, 0x40, 0x80)

subtitle = d.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run('操 作 手 册')
sub_run.font.size = Pt(28); sub_run.font.bold = True
sub_run.font.color.rgb = RGBColor(0x40, 0x95, 0xee)

d.add_paragraph()
info = d.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('\n\n\n\n文档版本:  V3.0\n更新日期:  2026-07-12\n适用版本:  v3.0+\n受 众:    客户 / 坐席 / 运营 / 管理员 / 运维').font.size = Pt(14)
d.add_page_break()

# 目录
d.add_heading('目录', level=1)
toc = [
    '第一篇  客户使用指南',
    '  第 1 章  注册与登录',
    '  第 2 章  开始咨询',
    '  第 3 章  消息与文件',
    '  第 4 章  转人工 / 评分',
    '第二篇  坐席工作指南',
    '  第 5 章  坐席登录与状态',
    '  第 6 章  接单与对话',
    '  第 7 章  模板回复 / 转接 / 标签',
    '  第 8 章  富媒体 (录像/语音/视频)',
    '第三篇  运营管理指南',
    '  第 9 章  数据看板',
    '  第 10 章  实时监控大屏',
    '  第 11 章  健康分与主动关怀',
    '第四篇  管理员配置',
    '  第 12 章  系统配置',
    '  第 13 章  限流 / 脱敏 / 告警',
    '  第 14 章  监控运维',
    '附录 A  常见问题 FAQ',
    '附录 B  故障排查',
    '附录 C  合规与安全',
    '附录 D  术语表',
]
for t in toc:
    d.add_paragraph(t)
d.add_page_break()

# ============ 第一篇 ============
d.add_heading('第一篇  客户使用指南', level=1)
d.add_paragraph('本篇面向使用 V3 智能客服平台的最终客户, 介绍如何注册、登录、咨询、转人工、评分。')

d.add_heading('第 1 章  注册与登录', level=2)
d.add_heading('1.1 演示账号 (快速体验)', level=3)
d.add_paragraph('V3 平台预置了 3 个客户账号, 可直接使用:')
t = d.add_table(rows=4, cols=3); t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '账号'; t.rows[0].cells[1].text = '密码'; t.rows[0].cells[2].text = '说明'
for i, (u, p, dd) in enumerate([
    ('customer1', '123456', '普通客户 (有历史会话)'),
    ('customer2', '123456', '普通客户'),
    ('customer3', '123456', 'VIP 客户 (优先接入)'),
]):
    t.rows[i+1].cells[0].text = u
    t.rows[i+1].cells[1].text = p
    t.rows[i+1].cells[2].text = dd

d.add_heading('1.2 首次访问', level=3)
for s in [
    '步骤 1: 打开浏览器 (推荐 Chrome / Edge / Safari 最新版), 访问:',
    '  生产环境: https://cs.example.com',
    '  测试环境: http://localhost (本地 Docker)',
    '步骤 2: 点击「立即注册」, 填写:',
    '  用户名 (4-20 字符, 字母+数字)',
    '  密码 (6+ 字符, 建议字母+数字+特殊字符)',
    '  昵称 (聊天时显示)',
    '  邮箱 (找回密码用)',
    '步骤 3: 注册成功后自动登录, 跳转到客服首页。',
]:
    d.add_paragraph(s)

d.add_heading('1.3 密码安全建议', level=3)
for s in [
    '  至少 12 位字符, 包含大小写+数字+特殊',
    '  定期更换 (建议 90 天)',
    '  不同网站用不同密码',
    '  不要把密码写在便签上',
    '  不要在公共电脑保存登录',
]:
    d.add_paragraph(s)

d.add_page_break()
d.add_heading('第 2 章  开始咨询', level=2)
d.add_heading('2.1 智能客服 vs 人工客服', level=3)
d.add_paragraph('V3 平台提供两种客服模式:')
t = d.add_table(rows=4, cols=3); t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '对比项'; t.rows[0].cells[1].text = '智能客服 (Bot)'; t.rows[0].cells[2].text = '人工客服 (Agent)'
for i, (a, b, c) in enumerate([
    ('响应速度', '< 100ms', '< 5s'),
    ('可用时间', '7x24', '工作时段'),
    ('回答深度', '常见问题', '复杂业务'),
]):
    t.rows[i+1].cells[0].text = a; t.rows[i+1].cells[1].text = b; t.rows[i+1].cells[2].text = c

d.add_heading('2.2 发起咨询', level=3)
for s in [
    '步骤 1: 登录后, 在首页点击「开始咨询」',
    '步骤 2: 选择技能 (general/billing/tech):',
    '  通用 (general): 常见问题',
    '  账单 (billing): 订单/支付',
    '  技术 (tech): 技术问题',
    '步骤 3: 系统自动接入智能客服, 开始对话',
    '步骤 4: 如需人工, 在输入框输入「人工」/「真人」/「转接」自动切换',
]:
    d.add_paragraph(s)

d.add_heading('2.3 快捷入口 (FAB 浮动按钮)', level=3)
d.add_paragraph('在屏幕右下角, 有 4 个浮动按钮:')
for icon, name, desc in [
    ('语音', '语音', '点击拨打 AI 智能电话, 语音对话'),
    ('扫一扫', '扫一扫', '扫描二维码 / 拍照识别'),
    ('社区', '社区', '浏览其他客户的问答'),
    ('视频', '视频', '一键视频咨询'),
]:
    d.add_paragraph(f'  {icon}  -  {desc}')

d.add_page_break()
d.add_heading('第 3 章  消息与文件', level=2)
d.add_heading('3.1 发送消息', level=3)
t = d.add_table(rows=6, cols=2); t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '操作'; t.rows[0].cells[1].text = '说明'
for i, (a, b) in enumerate([
    ('点击输入框', '光标进入, 即可输入文字'),
    ('点击表情按钮', '弹出 32 个常用表情, 点击插入'),
    ('点击录音按钮', '最长 60 秒, 自动转文字发送'),
    ('点击图片按钮', '选择本地图片 (<= 5MB)'),
    ('拖拽文件到聊天区', '松手即上传, 适用于图片/文档/压缩包'),
]):
    t.rows[i+1].cells[0].text = a; t.rows[i+1].cells[1].text = b

d.add_heading('3.2 键盘快捷键 (效率提升 30%)', level=3)
for k, dd in [
    ('Enter', '发送消息 (输入框内)'),
    ('Shift+Enter', '换行'),
    ('Esc', '关闭弹窗 / 抽屉'),
    ('Cmd/Ctrl+K', '打开消息搜索'),
    ('Cmd/Ctrl+/', '显示快捷键帮助'),
    ('Alt+C', '切换技能'),
    ('Alt+H', '申请转人工'),
]:
    d.add_paragraph(f'  {k}  -  {dd}')

d.add_heading('3.3 撤回消息', level=3)
d.add_paragraph('  在自己发送的消息上, 长按或右键 -> 「撤回」')
d.add_paragraph('  仅限 2 分钟内, 超时不允许撤回 (防止滥用)')
d.add_paragraph('  撤回后, 对方看到「对方撤回了一条消息」')

d.add_heading('3.4 草稿自动保存 (重要!)', level=3)
d.add_paragraph('  输入过程中, 内容自动保存到本地 (防丢失)')
d.add_paragraph('  切换会话 / 刷新页面 / 意外断网, 草稿不丢')
d.add_paragraph('  发送成功后, 草稿自动清空')

d.add_page_break()
d.add_heading('第 4 章  转人工 / 评分', level=2)
d.add_heading('4.1 申请转人工', level=3)
for s in [
    '方式 1: 文本触发 (推荐)',
    '  在聊天中输入以下任一关键词, 自动转人工:',
    '    「人工」/「真人」/「坐席」/「转接」/「转人工」/「agent」',
    '方式 2: 主动点击',
    '  在消息列表右上方, 点击「转人工」按钮 (如有)',
    '方式 3: 连续 3 次「无法理解」',
    '  Bot 连续 3 次无法回答, 自动弹窗提示转人工',
    '转人工后:',
    '  关闭当前 Bot 会话, 创建新人工会话',
    '  等待坐席接单 (一般 < 30s)',
    '  客户历史信息自动带过来, 无需重复描述',
]:
    d.add_paragraph(s)

d.add_heading('4.2 CSAT 评分', level=2)
d.add_paragraph('会话结束后, 弹出评分弹窗, 1-5 星:')
for a, b in [
    ('1', '非常不满意'),
    ('2', '不满意'),
    ('3', '一般'),
    ('4', '满意'),
    ('5', '非常满意'),
]:
    d.add_paragraph(f'  {a} 星  -  {b}')

d.add_paragraph('  评分后可以填写文字评论 (可选)')
d.add_paragraph('  提交后 1s 内自动关闭页面')
d.add_paragraph('  评分将用于健康分计算和坐席绩效')

d.add_page_break()
# ============ 第二篇 ============
d.add_heading('第二篇  坐席工作指南', level=1)
d.add_paragraph('本篇面向使用 V3 平台的客服坐席, 介绍如何登录、接单、回复、转接、使用富媒体。')

d.add_heading('第 5 章  坐席登录与状态', level=2)
d.add_heading('5.1 演示账号', level=3)
t = d.add_table(rows=4, cols=4); t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '账号'; t.rows[0].cells[1].text = '密码'; t.rows[0].cells[2].text = '姓名'; t.rows[0].cells[3].text = '技能'
for i, (u, p, s, sk) in enumerate([
    ('agent1', '123456', '客服-小张', 'general'),
    ('agent2', '123456', '客服-小李', 'billing'),
    ('agent3', '123456', '客服-小王', 'tech'),
]):
    t.rows[i+1].cells[0].text = u; t.rows[i+1].cells[1].text = p
    t.rows[i+1].cells[2].text = s; t.rows[i+1].cells[3].text = sk

d.add_heading('5.2 状态切换 (右上方下拉框)', level=3)
for s, desc in [
    ('在线 (ONLINE)', '可接新会话, 客户可见'),
    ('忙碌 (BUSY)', '暂停接新, 已有会话继续'),
    ('离开 (AWAY)', '短暂离开, 不会断开会话'),
    ('离线 (OFFLINE)', '彻底退出, 不接任何新会话'),
]:
    d.add_paragraph(f'  {s}  -  {desc}')

d.add_heading('5.3 桌面通知 (推荐开启)', level=3)
d.add_paragraph('  第一次进入坐席工作台, 浏览器会弹窗请求通知权限')
d.add_paragraph('  点击「允许」, 进线客户时桌面右下角会弹出通知')
d.add_paragraph('  点击通知, 自动跳到该客户并接单')
d.add_paragraph('  通知有 30s 冷却, 同一客户不重复弹')

d.add_page_break()
d.add_heading('第 6 章  接单与对话', level=2)
d.add_heading('6.1 接单方式', level=3)
for s in [
    '方式 1: 手动指定',
    '  左侧「等待列表」显示所有 WAITING 客户',
    '  每个客户有「接起」按钮, 点击即接 (防串线 CAS)',
    '  若别人先接了, 提示「已被 #X 接起, 请选其他」',
    '方式 2: 一键抢单',
    '  顶部「抢单」按钮, 自动从队列抢一个',
    '  适用于空闲时快速消化队列',
]:
    d.add_paragraph(s)

d.add_heading('6.2 防串线机制 (重要!)', level=3)
d.add_paragraph('V3 平台使用 MySQL CAS 防止多人同时抢同一会话:')
for s in [
    '  每个会话状态: WAITING -> ACTIVE -> CLOSED',
    '  多坐席同时接: 数据库行锁串行化',
    '  第一个 SQL UPDATE 成功的, agent_id 写入',
    '  其他坐席 affected=0, 收到 409 错误',
    '  提示友好: 「已被 #X 接起」',
    '  自动从队列移除该会话',
]:
    d.add_paragraph(s)

d.add_heading('6.3 主动关怀 / 智能回复', level=3)
d.add_paragraph('  输入时, 顶部显示「AI 智能建议」(基于历史 FAQ)')
d.add_paragraph('  点击「应用」直接填入输入框')
d.add_paragraph('  客户输入时, 实时显示情感分析 (正面 / 负面 / 中性)')

d.add_page_break()
d.add_heading('第 7 章  模板回复 / 转接 / 标签', level=2)
d.add_heading('7.1 模板回复 (快捷文本)', level=3)
for s in [
    '  输入框左侧「模板」按钮, 弹出模板选择',
    '  模板按技能分类 (general/billing/tech)',
    '  公共模板: 所有人可见',
    '  个人模板: 仅创建者可见',
    '  点击模板, 内容自动填入输入框',
]:
    d.add_paragraph(s)

d.add_heading('7.2 会话转接', level=3)
d.add_paragraph('  在会话窗口, 点击右上角「转接」图标')
d.add_paragraph('  选择目标坐席 (按技能过滤)')
d.add_paragraph('  填写转接原因 (如「账单问题, 需财务处理」)')
d.add_paragraph('  系统消息会通知客户「已转接给客服 XXX」')

d.add_heading('7.3 客户标签 / 健康分', level=3)
for s in [
    '  在客户头像旁, 显示健康分 (0-100)',
    '  颜色: 绿色 >= 80 健康, 黄色 50-79 关注, 红色 < 50 风险',
    '  点击健康分, 弹窗显示详细指标',
    '  可手动给客户打标签 (VIP / 投诉 / 高价值 / ...)',
]:
    d.add_paragraph(s)

d.add_page_break()
d.add_heading('第 8 章  富媒体 (录像 / 语音 / 视频)', level=2)
d.add_heading('8.1 合规录像 (重要!)', level=3)
for s in [
    '  首次开启, 弹出录制同意书 (合规要求)',
    '  客户和坐席均需同意才能继续',
    '  录像参数: 25 fps / 1080p / 2.5 Mbps (高清)',
    '  录制过程: 实时分片上传 -> 服务端 ffmpeg 合并',
    '  停止: 一键点击, 自动合并并保存到录像库',
    '  录像可在「回放」页面查看 (需管理员权限)',
]:
    d.add_paragraph(s)

d.add_heading('8.2 智能电话 (ASR + TTS)', level=3)
d.add_paragraph('  点击 FAB「语音」, 启动 WebRTC 通话')
d.add_paragraph('  客户说话 -> ASR 实时转文字 (显示在屏幕上)')
d.add_paragraph('  坐席回复 -> TTS 转语音播放给客户')
d.add_paragraph('  通话全程录音, 会话结束生成文字记录')

d.add_heading('8.3 视频通话 (WebRTC)', level=3)
d.add_paragraph('  客户和坐席均可发起, 需双方点击「接受」')
d.add_paragraph('  信令通过 STOMP /user/queue/video/{sid} 交换')
d.add_paragraph('  端到端 P2P 加密, 录像可保存到服务端')

d.add_page_break()
# ============ 第三篇 ============
d.add_heading('第三篇  运营管理指南', level=1)
d.add_paragraph('本篇面向运营 / 数据分析师, 介绍如何看数据、配置主动关怀、分析健康分。')

d.add_heading('第 9 章  数据看板', level=2)
d.add_heading('9.1 坐席数据看板 (AgentDashboard)', level=3)
d.add_paragraph('入口: 坐席工作台 -> 右上角「看板」')
d.add_paragraph('显示 6 大指标:')
for s, v in [
    ('今日会话数', '当前坐席今日接起的会话总数'),
    ('已解决率', '客户评分 >= 4 星的比例'),
    ('平均响应时间', '从客户发消息到坐席回复的平均时间'),
    ('平均会话时长', '从开始到结束的平均分钟数'),
    ('满意度 CSAT', '5 星制平均分'),
    ('当前活跃会话', '正在进行中 (ACTIVE) 的会话数'),
]:
    d.add_paragraph(f'  {s}  -  {v}')

d.add_heading('9.2 7 天趋势', level=3)
d.add_paragraph('  折线图, 显示会话量 / 满意度 / 响应时间 7 天趋势')
d.add_paragraph('  用于发现业务异常 (如周末激增 / 工作日下降)')

d.add_page_break()
d.add_heading('第 10 章  实时监控大屏', level=2)
d.add_heading('10.1 入口', level=3)
d.add_paragraph('  登录坐席账号 -> 头部「监控」按钮')
d.add_paragraph('  直接访问 /monitor 路由 (需坐席权限)')

d.add_heading('10.2 4 大核心指标', level=3)
t = d.add_table(rows=5, cols=2); t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '指标'; t.rows[0].cells[1].text = '说明'
for i, (a, b) in enumerate([
    ('今日总会话', '从 0 点到现在, 全平台所有会话数'),
    ('活跃会话', '状态 = ACTIVE (正在进行中)'),
    ('等候队列', '状态 = WAITING (待坐席接单)'),
    ('接通率', 'ACTIVE / (ACTIVE + WAITING + CLOSED)'),
]):
    t.rows[i+1].cells[0].text = a; t.rows[i+1].cells[1].text = b

d.add_heading('10.3 满意度分布', level=3)
d.add_paragraph('5 个进度条, 显示各星级比例:')
for s, p in [
    ('1 星 (非常不满意)', '红色, 一般 < 5%'),
    ('2 星 (不满意)', '橙色, 一般 < 10%'),
    ('3 星 (一般)', '灰色, 一般 20-30%'),
    ('4 星 (满意)', '浅蓝, 一般 30-40%'),
    ('5 星 (非常满意)', '深绿, 一般 30-40%'),
]:
    d.add_paragraph(f'  {s}  -  {p}')

d.add_heading('10.4 实时事件流', level=3)
d.add_paragraph('  大屏底部时间轴, 显示最近 20 个事件')
d.add_paragraph('  事件类型: SESSION_NEW / SESSION_ACTIVE / SESSION_CLOSED / RATED / TRANSFERRED')
d.add_paragraph('  实时推送 (STOMP), 0 延迟')

d.add_page_break()
d.add_heading('第 11 章  健康分与主动关怀', level=2)
d.add_heading('11.1 健康分算法', level=3)
d.add_paragraph('健康分 (0-100) 综合 6 个维度:')
t = d.add_table(rows=7, cols=3); t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '维度'; t.rows[0].cells[1].text = '权重'; t.rows[0].cells[2].text = '说明'
for i, (a, w, dd) in enumerate([
    ('满意度', '30%', '近 30 天 CSAT 平均'),
    ('响应速度', '20%', '平均响应时间 (越快越好)'),
    ('解决率', '20%', '已解决会话 / 总会话'),
    ('回访率', '10%', '7 天内再次咨询的比例'),
    ('活跃度', '10%', '近 30 天咨询频次'),
    ('消费贡献', '10%', 'VIP / 高价值客户权重'),
]):
    t.rows[i+1].cells[0].text = a; t.rows[i+1].cells[1].text = w; t.rows[i+1].cells[2].text = dd

d.add_heading('11.2 主动关怀规则 (10+ 模板)', level=3)
for s in [
    '  流失预警: 30 天未活跃 -> 推送优惠',
    '  高价值生日: 生日当天推送问候',
    '  新客引导: 注册 3 天未下单 -> 推送教程',
    '  投诉后回访: 评分 <= 2 星 -> 自动升级主管',
    '  节日关怀: 法定节日 / 双 11 等',
    '  升级 VIP: 累计消费 > 1000 -> 推送升级',
    '  续费提醒: 服务到期前 7 天',
    '  关联推荐: 浏览 X 推荐 Y',
]:
    d.add_paragraph(s)

d.add_page_break()
# ============ 第四篇 ============
d.add_heading('第四篇  管理员配置', level=1)
d.add_paragraph('本篇面向系统管理员 / SRE, 介绍系统配置、限流、监控、告警。')

d.add_heading('第 12 章  系统配置', level=2)
d.add_heading('12.1 配置文件 (application.yml)', level=3)
d.add_paragraph('主要配置项:')
t = d.add_table(rows=8, cols=3); t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '配置项'; t.rows[0].cells[1].text = '默认值'; t.rows[0].cells[2].text = '说明'
for i, (k, dv, dd) in enumerate([
    ('chat.jwt.secret', 'change-me-please-use-a-32-byte-secret', 'JWT 加密密钥, 必改!'),
    ('chat.jwt.ttl-ms', '86400000 (24h)', 'Token 过期时间'),
    ('spring.datasource.url', 'jdbc:mariadb://localhost:3306/cs_im', '数据库地址'),
    ('spring.redis.host', 'localhost:6379', 'Redis 地址'),
    ('chat.recording.enabled', 'true', '是否开启录像'),
    ('chat.ratelimit.login', '5 req/min', '登录限流'),
    ('chat.sla.response-target', '1000 ms', 'P99 响应时间 SLA'),
]):
    t.rows[i+1].cells[0].text = k; t.rows[i+1].cells[1].text = dv; t.rows[i+1].cells[2].text = dd

d.add_heading('12.2 环境变量 (.env)', level=3)
d.add_paragraph('Docker Compose 启动前, 配置 .env:')
for v, defv in [
    ('MYSQL_ROOT_PASSWORD', 'root 密码'),
    ('JWT_SECRET', 'JWT 密钥 (32 字节)'),
    ('LOG_LEVEL', '日志级别 (INFO/DEBUG)'),
    ('GRAFANA_ADMIN_PASSWORD', 'Grafana 管理员密码'),
]:
    d.add_paragraph(f'  {v}={defv or "(必填)"}')

d.add_page_break()
d.add_heading('第 13 章  限流 / 脱敏 / 告警', level=2)
d.add_heading('13.1 限流 (@RateLimit 注解)', level=3)
d.add_paragraph('在 Controller 方法上加注解:')
d.add_paragraph('  @RateLimit(key = "login", permits = 5, window = 60)')
d.add_paragraph('表示: 60 秒内最多 5 次登录请求, 超限返 429。')
d.add_paragraph('预置规则:')
for s in [
    '  登录: 5/min/IP',
    '  注册: 3/hour/IP',
    '  发送消息: 30/min/user',
    '  评分: 10/hour/user',
    '  抢单: 60/min/agent',
    '  全局默认: 100/min/endpoint',
]:
    d.add_paragraph(s)

d.add_heading('13.2 脱敏 (@Desensitize 注解)', level=3)
d.add_paragraph('自动识别并脱敏日志中的敏感字段:')
t = d.add_table(rows=6, cols=2); t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '字段'; t.rows[0].cells[1].text = '脱敏后'
for i, (a, b) in enumerate([
    ('手机号', '138****1234'),
    ('身份证', '110101********1234'),
    ('邮箱', 'li***@example.com'),
    ('银行卡', '6222 **** **** 1234'),
    ('姓名', '张* (单字) / 张*三 (复姓)'),
]):
    t.rows[i+1].cells[0].text = a; t.rows[i+1].cells[1].text = b

d.add_heading('13.3 告警通道 (4 个)', level=3)
t = d.add_table(rows=5, cols=3); t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '通道'; t.rows[0].cells[1].text = '适用'; t.rows[0].cells[2].text = '配置'
for i, (a, b, c) in enumerate([
    ('PagerDuty', '严重故障 (P0)', '5min 响应'),
    ('钉钉', '常规告警 (P1/P2)', '群机器人 + 通知所有人'),
    ('飞书', '常规告警 (P1/P2)', '群机器人 + @oncall'),
    ('Webhook', '自定义 (邮箱/Slack)', 'POST JSON 推送到指定 URL'),
]):
    t.rows[i+1].cells[0].text = a; t.rows[i+1].cells[1].text = b; t.rows[i+1].cells[2].text = c

d.add_page_break()
d.add_heading('第 14 章  监控运维', level=2)
d.add_heading('14.1 监控栈 (3 组件)', level=3)
t = d.add_table(rows=4, cols=3); t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '组件'; t.rows[0].cells[1].text = '端口'; t.rows[0].cells[2].text = '职责'
for i, (a, p, dd) in enumerate([
    ('Prometheus', 'localhost:9090', '时序数据库, 抓取 /actuator/prometheus'),
    ('AlertManager', 'localhost:9093', '告警评估 + 路由 + 抑制'),
    ('Grafana', 'localhost:3000', '可视化看板 (admin/admin)'),
]):
    t.rows[i+1].cells[0].text = a; t.rows[i+1].cells[1].text = p; t.rows[i+1].cells[2].text = dd

d.add_heading('14.2 12 条告警规则', level=3)
d.add_paragraph('预置告警规则 (ops/prometheus/alerts.yml):')
for k, cond in [
    ('P0_ErrorRate', '错误率 > 5% (5min)'),
    ('P0_ResponseP99', 'P99 响应时间 > 1s (5min)'),
    ('P0_JVMMemory', 'JVM 堆内存使用 > 90%'),
    ('P0_WebSocketDrop', 'WS 在线连接掉到 0 (5min)'),
    ('P1_DatabaseDown', 'MySQL 主库不可达'),
    ('P1_RedisDown', 'Redis 不可达'),
    ('P1_HighCPU', 'CPU > 80% (10min)'),
    ('P1_HighDisk', '磁盘使用 > 85%'),
    ('P2_LongGC', 'GC 时间 > 200ms'),
    ('P2_QueueBacklog', 'Redis 队列堆积 > 1000'),
    ('P2_CSATTlow', 'CSAT 平均 < 3.5 (1h)'),
    ('P2_AgentOffline', '所有坐席离线 > 30min'),
]:
    d.add_paragraph(f'  {k}  -  {cond}')

d.add_heading('14.3 常用运维命令', level=3)
d.add_paragraph('查看服务状态:')
d.add_paragraph('  docker-compose ps')
d.add_paragraph('查看日志 (实时):')
d.add_paragraph('  docker-compose logs -f cs-im')
d.add_paragraph('重启服务:')
d.add_paragraph('  docker-compose restart cs-im')
d.add_paragraph('查看 Prometheus 指标:')
d.add_paragraph('  curl http://localhost:9001/actuator/prometheus | grep http_server_requests_seconds')

d.add_page_break()
# ============ 附录 ============
d.add_heading('附录 A  常见问题 FAQ', level=1)
faq = [
    ('A.1 客户: 发送消息后没回复?', '可能原因: 1) 坐席全员离线 -> 等待上线; 2) 网络问题 -> 顶部黄色横幅; 3) 系统维护 -> 看首页公告。'),
    ('A.2 客户: 撤回失败?', '仅 2 分钟内可撤回, 超时不允许。'),
    ('A.3 客户: 草稿丢失?', '草稿自动保存到 localStorage, 清除浏览器数据会丢。建议重要内容先发出去。'),
    ('A.4 坐席: 抢单提示「已被 #X 接起」?', '防串线机制生效, 别人先接了, 请选其他会话。'),
    ('A.5 坐席: 录像卡顿?', '网络差会卡, 25fps/1080p 需要 >= 5Mbps 上行带宽。'),
    ('A.6 管理员: 大屏显示「暂无数据」?', '今天还没会话, 或后端 cs-success 不可达, 检查 /api/success/realtime 端点。'),
    ('A.7 管理员: Grafana 看板空白?', 'Prometheus 没抓到数据, 检查 docker-compose ps 看 cs-im 是否健康。'),
    ('A.8 管理员: 告警太多怎么办?', '调整 ops/prometheus/alerts.yml 阈值, 常用值: 错误率 5%->10%, P99 1s->2s。'),
    ('A.9 SRE: 数据库连接失败?', '检查 .env 的 MYSQL_ROOT_PASSWORD, 是否和 application.yml 一致。'),
    ('A.10 SRE: Redis 队列堆积?', '坐席不足, 临时增加坐席或加自动接单机器人。'),
]
for q, a in faq:
    d.add_heading(q, level=3)
    d.add_paragraph('答: ' + a)

d.add_page_break()
d.add_heading('附录 B  故障排查', level=1)
d.add_heading('B.1 故障等级', level=2)
t = d.add_table(rows=5, cols=4); t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '等级'; t.rows[0].cells[1].text = '响应'; t.rows[0].cells[2].text = '通知'; t.rows[0].cells[3].text = 'SLA'
for i, (a, b, c, dd) in enumerate([
    ('P0', '5 min', 'PagerDuty 全员', '5 min 修复 / 30 min 降级'),
    ('P1', '30 min', '钉钉/飞书群', '30 min 修复 / 4h 降级'),
    ('P2', '4h', '钉钉/飞书群', '1 天修复'),
    ('P3', '1 天', '邮件 / 工单', '1 周修复'),
]):
    t.rows[i+1].cells[0].text = a; t.rows[i+1].cells[1].text = b
    t.rows[i+1].cells[2].text = c; t.rows[i+1].cells[3].text = dd

d.add_heading('B.2 常见故障速查', level=2)
issues = [
    ('客户: 打不开页面', '1) 看 nginx 日志: docker-compose logs nginx; 2) 检查防火墙; 3) 看 index.html 是否 200'),
    ('客户: WebSocket 连不上', '1) 看 cs-gateway 日志; 2) 检查 STOMP 端点; 3) nginx 配置 proxy_set_header Upgrade $http_upgrade'),
    ('客户: 消息发送失败', '1) 看 Network 标签请求; 2) STOMP 断连会自动重连; 3) 看后端 logs/cs-im/error.log'),
    ('坐席: 接单 409', '防串线正常行为, 提示用户并刷新队列'),
    ('坐席: 录像上传失败', '1) 检查网络; 2) 后端磁盘: df -h; 3) 看 chat_record 表 status'),
    ('大屏: 数据不刷新', '1) STOMP 连上没: 看 console; 2) 5s 兜底轮询是否成功; 3) 浏览器 F12 查 /api/success/realtime'),
    ('监控: 告警风暴', '1) 临时禁告警: amtool silence add ...; 2) 调阈值; 3) 排查根因'),
    ('DB: 主库挂了', '1) 触发切换: docker-compose exec mysql-replica promote; 2) 修改 app 指向新主; 3) 复盘'),
]
for k, v in issues:
    d.add_heading(k, level=3)
    d.add_paragraph(v)

d.add_page_break()
d.add_heading('附录 C  合规与安全', level=1)
d.add_heading('C.1 PIPL / GDPR 合规', level=2)
for s in [
    '  用户数据加密存储 (AES-256)',
    '  敏感字段脱敏 (手机号/身份证/邮箱)',
    '  录像前需用户明示同意',
    '  数据保留 90 天后自动归档',
    '  用户可申请数据导出/删除 (GDPR 17 条)',
    '  所有操作审计日志 (audit_log 表)',
]:
    d.add_paragraph(s)

d.add_heading('C.2 安全建议', level=2)
for s in [
    '  JWT secret 必须 32+ 字节, 定期更换',
    '  数据库密码 16+ 字符, 包含大小写+数字+特殊',
    '  启用 HTTPS (生产环境必)',
    '  防火墙: 仅开放 80/443, 内部服务用 docker network',
    '  关闭 debug 模式 (spring.profiles.active=prod)',
    '  定期备份数据库 (cron 每日 0 点)',
]:
    d.add_paragraph(s)

d.add_heading('C.3 数据保留策略', level=2)
t = d.add_table(rows=5, cols=3); t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '数据类型'; t.rows[0].cells[1].text = '保留时长'; t.rows[0].cells[2].text = '处理'
for i, (a, b, c) in enumerate([
    ('聊天消息', '90 天', '到期后归档到 oss/cold'),
    ('录像文件', '180 天', '到期后删除 (可配置)'),
    ('审计日志', '365 天', '合规要求'),
    ('健康分历史', '永久', '用于趋势分析'),
]):
    t.rows[i+1].cells[0].text = a; t.rows[i+1].cells[1].text = b; t.rows[i+1].cells[2].text = c

d.add_page_break()
d.add_heading('附录 D  术语表', level=1)
terms = [
    ('CSAT', 'Customer Satisfaction Score, 客户满意度 (1-5 星)'),
    ('NPS', 'Net Promoter Score, 净推荐值'),
    ('Session', '客户与坐席的一次完整对话'),
    ('Ticket', '会话结束后的工单, 用于后续跟进'),
    ('SLA', 'Service Level Agreement, 服务等级协议'),
    ('P99', '99% 请求响应时间 (性能指标)'),
    ('CAS', 'Compare-And-Set, 乐观锁 (防串线)'),
    ('STOMP', 'Simple Text Oriented Messaging Protocol, WebSocket 子协议'),
    ('WebRTC', 'Web Real-Time Communication, 浏览器 P2P 通信'),
    ('CDP', 'Customer Data Platform, 客户数据平台'),
    ('FAQ', 'Frequently Asked Questions, 常见问题'),
    ('NLP', 'Natural Language Processing, 自然语言处理'),
    ('AOP', 'Aspect-Oriented Programming, 面向切面编程 (限流/脱敏)'),
    ('JWT', 'JSON Web Token, 无状态身份令牌'),
    ('PIPL', 'Personal Information Protection Law, 中国个人信息保护法'),
    ('GDPR', 'General Data Protection Regulation, 欧盟数据保护条例'),
]
for t, dd in terms:
    d.add_paragraph(f'  {t}  -  {dd}')

d.add_paragraph()
end = d.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
end.add_run('——— 文档结束 ———').font.size = Pt(12)
end.add_run('\n\n如有问题, 请联系: support@example.com')
end.add_run('\nV3 Team · 2026-07-12')

import os
out = '/workspace/online-chat/docs/OPERATION-MANUAL.docx'
d.save(out)
print(f'+ {out}: {os.path.getsize(out)} bytes ({os.path.getsize(out)/1024:.1f} KB)')
print(f'+ 段落: {len(d.paragraphs)}, 表格: {len(d.tables)}')
